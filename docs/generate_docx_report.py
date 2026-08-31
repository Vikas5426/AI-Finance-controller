import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def add_code_block(doc, code_text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_background(cell, "F4F6F9")
    set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(code_text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x24, 0x29, 0x2E)
    doc.add_paragraph() # Spacing

def build_word_document(output_path):
    doc = Document()
    
    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Styles Setup
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # -------------------------------------------------------------
    # Document Title / Cover Section
    # -------------------------------------------------------------
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(10)
    title_p.paragraph_format.space_after = Pt(4)
    run_title = title_p.add_run("Autonomous AI Financial Controller")
    run_title.font.name = "Calibri"
    run_title.font.size = Pt(24)
    run_title.bold = True
    run_title.font.color.rgb = RGBColor(0x0F, 0x29, 0x42) # Deep Navy

    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_before = Pt(0)
    sub_p.paragraph_format.space_after = Pt(16)
    run_sub = sub_p.add_run("Core Logic, Mathematical Reconciler & Architecture Report")
    run_sub.font.name = "Calibri"
    run_sub.font.size = Pt(14)
    run_sub.font.color.rgb = RGBColor(0x00, 0x7A, 0xFF) # Vibrant Blue

    meta_table = doc.add_table(rows=2, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell_meta = meta_table.cell(0, 0)
    cell_meta.text = "System: AI Financial Controller v1.0.0"
    cell_meta.paragraphs[0].runs[0].font.size = Pt(9.5)
    cell_meta.paragraphs[0].runs[0].font.bold = True
    
    cell_meta2 = meta_table.cell(0, 1)
    cell_meta2.text = "Scope: 3-Way Reconciliation (Gateway ↔ Bank ↔ Ledger)"
    cell_meta2.paragraphs[0].runs[0].font.size = Pt(9.5)
    
    cell_meta3 = meta_table.cell(1, 0)
    cell_meta3.text = "Date: March 2026 / Active Environment"
    cell_meta3.paragraphs[0].runs[0].font.size = Pt(9.5)
    
    cell_meta4 = meta_table.cell(1, 1)
    cell_meta4.text = "Security: SHA-256 Audit Chain & Maker-Checker Enabled"
    cell_meta4.paragraphs[0].runs[0].font.size = Pt(9.5)

    for row in meta_table.rows:
        for cell in row.cells:
            set_cell_background(cell, "F0F4F8")
            set_cell_margins(cell, top=80, bottom=80, left=120, right=120)

    doc.add_paragraph()

    # -------------------------------------------------------------
    # 1. Executive Summary & Architecture Topology
    # -------------------------------------------------------------
    h1 = doc.add_heading(level=1)
    r1 = h1.add_run("1. Executive Summary & Architecture Topology")
    r1.font.color.rgb = RGBColor(0x0F, 0x29, 0x42)

    p = doc.add_paragraph(
        "The Autonomous AI Financial Controller is an enterprise financial intelligence platform that performs "
        "autonomous, high-throughput, three-way settlement reconciliation across Payment Gateways (e.g. Razorpay, Stripe), "
        "Bank Statements (e.g. HDFC/ICICI MT940 or CSV feeds), and General Ledger ERP entries (NetSuite, SAP, Tally)."
    )
    doc.add_paragraph(
        "Unlike generic LLM wrappers, this system is built on a quality-first deterministic foundation: it pairs mathematical "
        "matching algorithms (Hungarian assignment, dynamic programming subset-sum solvers, and RapidFuzz scoring) with a "
        "strictly bounded AI agent runtime that intercepts model reasoning behind an arithmetic verifier gate."
    )

    doc.add_paragraph("Key Architectural Pillars:").runs[0].bold = True
    pillars = [
        ("Multi-Pass Reconciler (P0–P5):", " Deduplication, 1:1 exact reference indices, N:1 dynamic programming netting, and global Hungarian fuzzy matching."),
        ("Deterministic Verifier Gate:", " LLM proposals are rejected if claimed fee/tax breakdowns do not mathematically match actual transaction variances."),
        ("4-Tier Decision Routing:", " Resolves clean transactions deterministically while isolating genuine boundary timing lags and missing credits."),
        ("Cryptographic Audit Trail:", " SHA-256 blockchain-style hash chain ensuring immutable, tamper-evident regulatory records."),
        ("Maker-Checker Authorization:", " Analyst-proposed adjusting journal entries require controller check-off before ledger commit.")
    ]
    for title, desc in pillars:
        bp = doc.add_paragraph(style='List Bullet')
        r_t = bp.add_run(title)
        r_t.bold = True
        bp.add_run(desc)

    # -------------------------------------------------------------
    # 2. Directory Map & File Locations
    # -------------------------------------------------------------
    h2 = doc.add_heading(level=1)
    r2 = h2.add_run("2. Directory Map & File Locations")
    r2.font.color.rgb = RGBColor(0x0F, 0x29, 0x42)

    doc.add_paragraph("Every module in the codebase has a distinct separation of concerns:")

    file_matrix = [
        ("API Entry Point", "backend/app/main.py", "FastAPI lifecycle, router registration, CORS, and static UI file mounting."),
        ("Batch Runner & Demo", "run_demo.py", "End-to-end demo execution: dataset generation, 6-pass matching, benchmarks, web server boot."),
        ("Matching Engine", "backend/app/services/matching_engine.py", "Core 6-pass reconciliation engine (P0 dedupe, P1 exact, P3 Hungarian fuzzy, P4 subset-sum DP, P5 residuals)."),
        ("Batch Orchestrator", "backend/app/services/batch_orchestrator.py", "Controlled 20–30 record analysis windows, 7-stage processing, and progress state machine."),
        ("Context Builder", "backend/app/services/context_builder.py", "360° financial contextualizer (historical fee profile, T+2 lag detection, cross-source candidate links)."),
        ("AI Agent Runtime", "backend/app/services/agent_runtime.py", "Bounded LLM investigation engine (Gemini/Claude) with deterministic arithmetic verifier gates."),
        ("Decision Engine", "backend/app/services/decision_engine.py", "4-tier decision routing (RESOLVED, RESOLVED_WITH_EXPLANATION, NEEDS_REVIEW, UNRESOLVED_EXCEPTION)."),
        ("Rules Engine", "backend/app/services/rules_engine.py", "Zero-dependency declarative rule evaluator for SOP thresholds, fee allowances, and write-offs."),
        ("Audit Hash Chain", "backend/app/services/audit_chain.py", "SHA-256 tamper-evident cryptographic hash chain recording all financial state mutations."),
        ("Cash Forecaster", "backend/app/services/cash_forecaster.py", "13-week forward liquidity runway segmented into Confirmed, Probable, At-Risk, and Unknown buckets."),
        ("Normalizer Service", "backend/app/services/normalizer.py", "Regular-expression reference key extractor, text tokenization, and canonical model mapper."),
        ("Data Validation", "backend/app/services/validation_service.py", "Structural sanity checks (nulls, negative amounts, future dates, ISO currencies)."),
        ("Synthetic Generator", "backend/app/services/synthetic_generator.py", "Generates 240+ multi-source records modeling 12 realistic financial topologies with ground truth."),
        ("Benchmarks", "backend/app/services/benchmarks.py", "Computes Precision, Recall, F1-Score, and Expected Calibration Error (ECE) against ground truth."),
        ("Schemas & Models", "backend/app/models/schemas.py", "Pydantic v2 schemas: CanonicalTransaction, MatchSchema, ExceptionSchema, ReconciliationDecision."),
        ("Database Schema", "backend/app/db/schema.py", "SQLAlchemy relational tables with foreign keys, compound indexes, and JSON payloads."),
        ("Database Service", "backend/app/db/database_service.py", "Database repository layer, transaction commits, seed users with Argon2 hashes."),
        ("Conversational Q&A", "backend/app/api/v1/qa.py", "Dynamic Insight Cards, visual timeline steps, and evidence checklists for the UI financial analyst."),
        ("Approvals API", "backend/app/api/v1/approvals.py", "Maker-checker dual-authorization workflows (propose, approve, reject with audit hashes)."),
        ("Batches API", "backend/app/api/v1/batches.py", "Batch creation, windowed reconciliation trigger, and progress streaming."),
        ("Frontend UI", "frontend/index.html", "Single-page console with window-by-window visualizer, match breakdown, and Q&A chat.")
    ]

    tbl_files = doc.add_table(rows=1, cols=3)
    tbl_files.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = tbl_files.rows[0].cells
    hdr_cells[0].text = "Component"
    hdr_cells[1].text = "File Path"
    hdr_cells[2].text = "Responsibility"
    for cell in hdr_cells:
        set_cell_background(cell, "0F2942")
        p = cell.paragraphs[0]
        p.runs[0].font.bold = True
        p.runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.runs[0].font.size = Pt(9.5)
        set_cell_margins(cell, top=100, bottom=100, left=120, right=120)

    for comp, path, resp in file_matrix:
        row_cells = tbl_files.add_row().cells
        row_cells[0].text = comp
        row_cells[0].paragraphs[0].runs[0].font.bold = True
        row_cells[0].paragraphs[0].runs[0].font.size = Pt(9)
        
        row_cells[1].text = path
        row_cells[1].paragraphs[0].runs[0].font.name = "Consolas"
        row_cells[1].paragraphs[0].runs[0].font.size = Pt(8.5)
        
        row_cells[2].text = resp
        row_cells[2].paragraphs[0].runs[0].font.size = Pt(9)
        
        for cell in row_cells:
            set_cell_margins(cell, top=60, bottom=60, left=100, right=100)

    doc.add_paragraph()

    # -------------------------------------------------------------
    # 3. Core Logic & Mathematical Workflows
    # -------------------------------------------------------------
    h3 = doc.add_heading(level=1)
    r3 = h3.add_run("3. Core Logic & Mathematical Workflows")
    r3.font.color.rgb = RGBColor(0x0F, 0x29, 0x42)

    # 3.1 Normalizer
    doc.add_heading("3.1 Ingestion, Decimal Precision & Reference Key Regex", level=2)
    doc.add_paragraph(
        "To eliminate floating-point arithmetic errors inherent in standard IEEE 754 representations, all amounts "
        "are ingested and normalized into integer minor units (paise for INR, cents for USD). For example, ₹10,000.00 is strictly "
        "handled as 1,000,000 paise throughout memory and database storage."
    )
    doc.add_paragraph("Regular expressions in NormalizerService extract typed reference keys:")
    add_code_block(doc, 
        "# Invoice: INV-YYYY-NNNN\n"
        "inv_matches = re.findall(r'INV-\\d{4}-\\d{4}', text, re.IGNORECASE)\n"
        "# Payment Key: pay_LtPk29Xq7\n"
        "pay_matches = re.findall(r'pay_[A-Za-z0-9]{6,}', text)\n"
        "# Settlement Key: SETL9KA22\n"
        "setl_matches = re.findall(r'setl_?[A-Za-z0-9]+', text, re.IGNORECASE)\n"
        "# Bank UTR Reference: N2604029912 / R2604029912\n"
        "utr_matches = re.findall(r'[NR]\\d{10,}', text, re.IGNORECASE)"
    )

    # 3.2 6-Pass Engine
    doc.add_heading("3.2 The 6-Pass Matching Pipeline", level=2)
    doc.add_paragraph(
        "The ReconciliationEngine executes a sequential multi-pass strategy. Reconciled records are locked and excluded from downstream passes:"
    )
    
    passes = [
        ("Pass P0 (Deduplication):", " Intra-source duplicate check over (source_kind, external_id) tuples. Flagged duplicate records are routed to Low-Severity exception queues."),
        ("Pass P1 (Exact 1:1 Keys):", " Inverted index lookup over payment, UTR, and invoice references. Confirms exact amount equality and records matches with 0.999 confidence."),
        ("Pass P4 (N:1 Settlement DP):", " Bounded Dynamic Programming knapsack solver that matches aggregated bank settlement credits against N individual gateway payments net of 2.0% MDR + 18% GST."),
        ("Pass P3 (Hungarian Fuzzy):", " Computes weighted feature vectors (ID, Amount, Date, Description, Counterparty, Account Code) and executes global linear sum assignment with a runner-up margin threshold >= 0.05."),
        ("Pass P5 (Residual Taxonomy):", " Classifies all remaining unmatched items into 16 standardized controller exception categories.")
    ]
    for p_title, p_desc in passes:
        bp = doc.add_paragraph(style='List Bullet')
        bp.add_run(p_title).bold = True
        bp.add_run(p_desc)

    # 3.3 Dynamic Programming Settlement Formulation
    doc.add_heading("3.3 N:1 Settlement Solver Formulation", level=2)
    doc.add_paragraph(
        "When a bank receives a lump-sum payout of target paise T from a gateway without individual itemized keys, "
        "the DP solver quantizes candidates by q = 100 and computes reachable sum combinations within a +/-3 day window:"
    )
    add_code_block(doc,
        "T = target_amount // 100\n"
        "reach: Dict[int, Tuple[int, ...]] = {0: ()}\n"
        "for idx, c in enumerate(cands[:60]):\n"
        "    fee = round(c.amount_minor * 0.02)\n"
        "    tax = round(fee * 0.18)\n"
        "    net_paise = c.amount_minor - fee - tax\n"
        "    vq = net_paise // 100\n"
        "    for s, chosen in list(reach.items()):\n"
        "        ns = s + vq\n"
        "        if ns <= T + 2 and ns not in reach:\n"
        "            reach[ns] = chosen + (idx,)\n"
        "best = min((s for s in reach if abs(s - T) * 100 <= 100), key=lambda s: abs(s - T), default=None)"
    )

    # 3.4 Bounded AI Runtime & Verifier Gate
    doc.add_heading("3.4 Deterministic Verifier Gate for LLM Reasoning", level=2)
    doc.add_paragraph(
        "To enforce 100% mathematical auditability, the AIAgentRuntime executes a DeterministicVerifier on every LLM proposal. "
        "If a model suggests an adjustment due to gateway processing fees, the verifier validates that the claimed fee breakup sum "
        "is identical to the actual minor variance:"
    )
    add_code_block(doc,
        "class DeterministicVerifier:\n"
        "    @staticmethod\n"
        "    def verify_proposal(proposal, exception_ctx, valid_txn_ids):\n"
        "        for cand_id in proposal.candidate_match_ids:\n"
        "            if cand_id not in valid_txn_ids:\n"
        "                return False, f'Candidate {cand_id} does not exist in batch.'\n"
        "        if proposal.recommended_action == 'ADJUST_LEDGER_FEE_SPLIT':\n"
        "            actual_diff = exception_ctx.get('impact_minor', 0)\n"
        "            claimed_sum = sum(int(v) for v in fee_breakup.values())\n"
        "            if claimed_sum != actual_diff:\n"
        "                return False, f'Arithmetic mismatch: {claimed_sum} != {actual_diff}.'\n"
        "        return True, None"
    )

    # 3.5 Hybrid Decision Tiers
    doc.add_heading("3.5 4-Tier Operational Routing Policy", level=2)
    tiers = [
        ("Tier 1 (RESOLVED):", " 100% Deterministic match with exact reference key and identical amount. Auto-closed with 1.00 confidence."),
        ("Tier 2 (RESOLVED_WITH_EXPLANATION):", " Contextual match verified against MDR schedules (2.0% or 1.5% + GST) within +/-₹1.00 rounding. Auto-closed with ~0.92 confidence."),
        ("Tier 3 (NEEDS_REVIEW):", " Period boundary cutoff or timing lag (T+2 settlement window). Generates an Analyst proposal to accrue to Account 1290 (In-Transit Clearing). Requires Maker-Checker check-off."),
        ("Tier 4 (UNRESOLVED_EXCEPTION):", " Unmatched bank credits, missing wires, or withheld chargebacks. Categorized as High/Critical severity for senior controller intervention.")
    ]
    for t_title, t_desc in tiers:
        bp = doc.add_paragraph(style='List Bullet')
        bp.add_run(t_title).bold = True
        bp.add_run(t_desc)

    # 3.6 Cryptographic Audit Hash Chain
    doc.add_heading("3.6 SHA-256 Cryptographic Audit Hash Chain", level=2)
    doc.add_paragraph(
        "Every event generates an immutable SHA-256 block cryptographically linked to the previous block's hash. "
        "Genesis starts with 64 zeros. Any manual manipulation of database records breaks chain verification immediately:"
    )
    add_code_block(doc,
        "preimage = f'{prev_hash}|{org_id}|{event_seq}|{event_type}|{entity_id}|{actor_id}|{ts_str}|{canonical_json_payload}'\n"
        "event_hash = hashlib.sha256(preimage.encode('utf-8')).hexdigest()"
    )

    # -------------------------------------------------------------
    # 4. Detailed Concrete Examples
    # -------------------------------------------------------------
    h4 = doc.add_heading(level=1)
    r4 = h4.add_run("4. Concrete Step-by-Step Code Examples")
    r4.font.color.rgb = RGBColor(0x0F, 0x29, 0x42)

    doc.add_heading("Example 1: Resolving a ₹10,000.00 Payment with 2% MDR Fee (pay_1002)", level=2)
    doc.add_paragraph("Scenario: Razorpay captured ₹10,000.00, but the bank statement shows a net credit of ₹9,764.00.")
    doc.add_paragraph("1. Raw Inputs:").runs[0].bold = True
    doc.add_paragraph("   • Gateway: Amount = 1,000,000 paise (₹10,000.00), Fee = 20,000 paise, Tax = 3,600 paise.\n   • Bank Line: Credit = 9,764.00 (976,400 paise), Ref = pay_1002.")
    doc.add_paragraph("2. Mathematical Resolution:").runs[0].bold = True
    doc.add_paragraph("   • Variance = 1,000,000 - 976,400 = 23,600 paise (₹236.00).\n   • Expected Fee = 1,000,000 * 0.02 = 20,000 paise (₹200.00).\n   • Expected GST = 20,000 * 0.18 = 3,600 paise (₹36.00).\n   • Total Net Fee = 23,600 paise. Variance exactly equals expected fee.")
    doc.add_paragraph("3. Accounting Result:").runs[0].bold = True
    doc.add_paragraph("   • Decision Tier: RESOLVED_WITH_EXPLANATION (Confidence: 0.92).\n   • Ledger Adjustment: Debit Account 5010 (Processing Fees) ₹236.00, Credit Account 1210 (Receivable) ₹236.00.")

    doc.add_heading("Example 2: Period Boundary Cutoff (T+2 Settlement Lag)", level=2)
    doc.add_paragraph("Scenario: Transaction captured on March 31, 2026 at 23:55:00 IST. Bank statement closes on March 31 without settlement credit.")
    doc.add_paragraph("1. Context Builder Flags: PERIOD_BOUNDARY_CUTOFF_T2_LAG, is_period_cutoff = True.")
    doc.add_paragraph("2. Hybrid Decision: Routed to Tier 3 (NEEDS_REVIEW) with 0.88 confidence.")
    doc.add_paragraph("3. Maker-Checker Card: Analyst proposes accrual entry: Debit 1290 (In-Transit Clearing) / Credit 4000 (Revenue). Controller Checker signs off at /api/v1/approvals/{id}/approve.")

    # -------------------------------------------------------------
    # 5. Database Schema & REST API Reference
    # -------------------------------------------------------------
    h5 = doc.add_heading(level=1)
    r5 = h5.add_run("5. Database Schema & API Reference")
    r5.font.color.rgb = RGBColor(0x0F, 0x29, 0x42)

    doc.add_heading("5.1 Relational Tables in SQLite / PostgreSQL", level=2)
    tables_list = [
        ("organizations", "Multi-tenant tenant boundaries, default base currency, materiality threshold."),
        ("users", "Controller user accounts, Argon2 password hashes, role-based limits (analyst, approver, admin)."),
        ("source_profiles", "Column mapping dictionaries and datetime format parser configurations."),
        ("batches", "Batch processing runs, total records, window counts, match rates, and execution time."),
        ("transactions", "Canonical transaction store with typed reference keys, minor amounts, and status."),
        ("matches", "Reconciled match entities (1:1, 1:N, N:1) with solver evidence JSON and confidence."),
        ("match_legs", "Individual legs of a match (primary vs. counterpart with signed minor amounts)."),
        ("exception_records", "16-type taxonomy exceptions with severity, impact amount, and AI investigation JSON."),
        ("approval_requests", "Maker-checker review queue with analyst proposals and approver sign-offs."),
        ("audit_events", "Immutable SHA-256 hash chain blocks with prev_hash, event_hash, and payload."),
        ("rules", "Declarative business rules for SOP fee thresholds and automated write-offs.")
    ]
    for t_name, t_desc in tables_list:
        bp = doc.add_paragraph(style='List Bullet')
        bp.add_run(t_name).bold = True
        bp.add_run(f": {t_desc}")

    doc.add_heading("5.2 REST API v1 Endpoints", level=2)
    api_list = [
        ("POST /api/v1/batches/create", "Initializes a new reconciliation batch and ingests multi-source data."),
        ("POST /api/v1/batches/run-windowed-pipeline", "Executes the 24-record windowed pipeline with progress state streaming."),
        ("GET /api/v1/transactions", "Filter and search canonical transactions across source, status, or reference keys."),
        ("GET /api/v1/exceptions", "Lists detected exceptions with severity and recommended actions."),
        ("POST /api/v1/exceptions/{id}/investigate", "Triggers deep AI agent investigation over a specific exception."),
        ("GET /api/v1/approvals/proposals", "Fetches pending Maker-Checker cards for analyst/approver review."),
        ("POST /api/v1/approvals/{id}/approve", "Approver checker sign-off, committing ledger adjustment & audit block."),
        ("GET /api/v1/audit/verify-chain", "Sequentially re-computes and verifies the complete SHA-256 cryptographic audit chain."),
        ("GET /api/v1/reports/cash-forecast-13w", "Generates the 13-week segmented liquidity runway forecast."),
        ("POST /api/v1/qa/ask", "Progressive-disclosure conversational finance assistant with dynamic insight cards.")
    ]
    for ep, ep_desc in api_list:
        bp = doc.add_paragraph(style='List Bullet')
        r_ep = bp.add_run(ep)
        r_ep.font.name = "Consolas"
        r_ep.font.size = Pt(9.5)
        r_ep.bold = True
        bp.add_run(f" — {ep_desc}")

    # Save document
    doc.save(output_path)
    print(f"[+] Successfully generated Word Document at: {output_path}")

if __name__ == "__main__":
    output_file = r"d:\AI Finance controller\docs\AI_Financial_Controller_Core_Logic_Report.docx"
    build_word_document(output_file)
    # Also save a copy in the root folder for direct access
    root_file = r"d:\AI Finance controller\AI_Financial_Controller_Core_Logic_Report.docx"
    build_word_document(root_file)
